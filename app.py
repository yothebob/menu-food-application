# food program
import random
from docx import Document
from datetime import date
import ast
import re

"""
    This is a food menu program designed to create cooking menus, you can store recipes, add ingredients and more!

"""



def import_dinner_dict():
    '''
    Input : NA
    Output : res (Dictionary)

    Imports a dictionary of meals from data/dinners_dict.txt
    '''
    file = open("data/dinners_dict.txt", "r")
    contents = file.read()
    res = ast.literal_eval(contents)
    file.close()
    return res



def import_ingredient_profiles():
    '''
    input : NA
    Output : res (dictionary)

    Imports a dictionary of ingredient flavor vectors from data/ingredients.txt
    '''
    f = open('data/ingredients.txt','r')
    contents = f.read()
    res = ast.literal_eval(contents)
    f.close()
    return res


#initializing
dinners = import_dinner_dict()
days = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
ingredient_profiles = import_ingredient_profiles()



def create_menu():
    '''
    Input : NA
    Output : menu (list)

    Creates a list of randomly picked meals, one for each day in the days
    '''
    return [random.choice(list(dinners.keys())) for day in days]



def add_meal():
    '''
    Input : NA
    Output : NA

    a function to manually add a meal/ingredients to "data/dinners_dict.txt"
    '''
    new_meal = input("what would you like to add?\n : ")
    ingredients = input("what ingredients do you use? (Please list as such ...,...,...)\n : ")

    ingredient_array = ingredients.split(",")
    print(ingredient_array)
    dinners[new_meal] = ingredient_array

    f = open("data/dinners_dict.txt", "w")
    f.write(str(dinners))
    f.close()



def generate_menu():
    '''
    Input : NA
    Output : NA

    a function to create your own weekly menu or randomly generate one.
    '''
    user_input = input("Do you want to create your own menu? (y/n)\n: ")

    if "y" in user_input.lower():
        menu = []
        for day in days:
            print(dinners.keys())
            user_input = input(f"type the recipe you want for {day}\n:")
            while user_input.lower() not in dinners.keys():
                print(dinners.keys())
                user_input = input(f"type the recipe you want for {day}\n:")
            menu.append(user_input)
    else:
        menu = create_menu()

    [print(days[x]+ ': ' + menu[x]) for x in range(len(menu))]
    save = input('does this menu look good to you? (y/n) to save.')

    if save.lower() == "y":
        f = open('data/saved_menus.txt','a')
        f.write(str(menu) +'|' + str(date.today()) + "\n")
        f.close()



def read_saved_menus():
    #read from saved_menus.txt, returns dict in format {date-index: menu}
    f = open('data/saved_menus.txt','r')
    saved_menus = {}
    index = 0

    for line in f:
        menu=[]
        seperate = line.split("|")
        if len(seperate) > 1:
            if "\n" in seperate[1]:
                date = seperate[1][:-2] + f"-{index}"
            else:
                date = seperate[1] + f"{index}"
            index += 1
            m = seperate[0][1:-1]
            mm = m.split(',')
            for item in mm:
                menu.append(item.strip("' "))
            saved_menus[date] = menu
    return saved_menus


def pick_saved_menu():
    
    menus = read_saved_menus()
    menu_index = []
    i = 0

    for key,val in menus.items():
        print(i,key,val)
        menu_index.append(key)
        i += 1

    pick_meal = input('Which meal would you like to use? pick corrisponding number (0-x) \n: ')
    menu = menus[menu_index[int(pick_meal)]]
    return menu



def create_docx():
    #create word doc (.docx) menu, generate new menu or use a saved menu you have
    new_or_saved = input('would you like to use a saved menu? (y/n) \n: ')

    if "y" in new_or_saved.lower():
        menu = pick_saved_menu()
    else:
        menu = create_menu()

    doc = Document()
    p = doc.add_paragraph()

    for day in range(6):
        p.add_run( days[day] + ": " + menu[day] + "\n")

    doc.save(str(date.today()) + "menu.docx")
    f = open("data/menus.txt","a")
    f.write(str(date.today()) + str(menu))
    f.close()



def create_grocery_list(menu):
    grocery_list = []
    for meal in menu:
        for ingredient in dinners[meal]:
            grocery_list.append(ingredient)
    print(grocery_list)
    return grocery_list



def terminal_interface():
    print('''
    Welcome to the Food Menu Program!!!
        type "menu" to create a menu
        type "docx" to create a word doc for a menu
        type "groceries" to create grocery list
        type "exit" to exit

    ''')
    user_input = input(": ")

    if user_input.lower() == "menu":
        generate_menu()
    elif user_input.lower() == "docx":
        create_docx()
    elif user_input.lower() == "groceries":
        create_grocery_list(pick_saved_menu())


''' irrelivent stuff below '''
#flavor profile for meal ingredients
flavor_profile = ['sweet','sour','salt','bitter','acidic','basic','savory','hotness','spiciness','oily','minty'
,'astringent','starchiness','horseradish','creamy','earthy']



def meal_to_vec(meal):
    #converting ingredients to a total meal flavor profile
    res_profile = [0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0]
    profiles = []
    for ingredient in dinners[meal]:
        if ingredient in ingredient_profiles.keys():
            print(ingredient,ingredient_profiles[ingredient])
            profiles.append(ingredient_profiles[ingredient])
        else:
            new_ingredient_profile = []
            print(ingredient)
            for flavor in range(len(flavor_profile)):
                profile = input(f"how {flavor_profile[flavor]} is this? (0-1)")
                new_ingredient_profile.append(float(profile))
            ingredient_profiles[ingredient] = new_ingredient_profile
            f = open('data/ingredient.txt','w')
            f.write(ingredient_profiles)
            f.close()
    for item in range(len(profiles)):
        for flavor in range(len(res_profile)):
            res_profile[flavor] += profiles[item][flavor]
    for item in res_profile:
        item = item/len(profiles)
    print(res_profile)

#meal_to_vec("testing")
#print(create_menu())
